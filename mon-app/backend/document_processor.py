# document_processor.py - Module de traitement des documents
import os
import zipfile
from typing import Optional
import docx
import PyPDF2
import pdfplumber
from PIL import Image
import pytesseract
import magic

class DocumentProcessor:
    """Classe pour traiter différents types de documents"""
    
    def __init__(self):
        """Initialiser le processeur de documents"""
        self.supported_formats = {'.txt', '.docx', '.pdf', '.doc'}
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        
    def extract_text(self, filepath: str) -> Optional[str]:
        """Extraire le texte d'un fichier selon son format"""
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Fichier non trouvé: {filepath}")
        
        # Vérifier la taille du fichier
        if os.path.getsize(filepath) > self.max_file_size:
            raise ValueError("Fichier trop volumineux")
        
        # Déterminer le type de fichier
        file_extension = os.path.splitext(filepath)[1].lower()
        
        try:
            if file_extension == '.txt':
                return self._extract_from_txt(filepath)
            elif file_extension == '.docx':
                return self._extract_from_docx(filepath)
            elif file_extension == '.pdf':
                return self._extract_from_pdf(filepath)
            elif file_extension == '.doc':
                return self._extract_from_doc(filepath)
            else:
                raise ValueError(f"Format non supporté: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction: {str(e)}")
    
    def _extract_from_txt(self, filepath: str) -> str:
        """Extraire le texte d'un fichier TXT"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as file:
                    content = file.read()
                    return self._clean_text(content)
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Impossible de décoder le fichier texte")
    
    def _extract_from_docx(self, filepath: str) -> str:
        """Extraire le texte d'un fichier DOCX"""
        try:
            doc = docx.Document(filepath)
            full_text = []
            
            # Extraire le texte des paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extraire le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            content = '\n'.join(full_text)
            return self._clean_text(content)
            
        except Exception as e:
            raise Exception(f"Erreur lors de la lecture du fichier DOCX: {str(e)}")
    
    def _extract_from_pdf(self, filepath: str) -> str:
        """Extraire le texte d'un fichier PDF"""
        text_content = ""
        
        # Méthode 1: PDFplumber (meilleure pour la mise en forme)
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            if text_content.strip():
                return self._clean_text(text_content)
        except Exception as e:
            print(f"Erreur avec pdfplumber: {e}")
        
        # Méthode 2: PyPDF2 (fallback)
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            if text_content.strip():
                return self._clean_text(text_content)
        except Exception as e:
            print(f"Erreur avec PyPDF2: {e}")
        
        # Méthode 3: OCR avec Tesseract (pour PDFs scannés)
        try:
            text_content = self._extract_with_ocr(filepath)
            if text_content.strip():
                return self._clean_text(text_content)
        except Exception as e:
            print(f"Erreur avec OCR: {e}")
        
        raise Exception("Impossible d'extraire le texte du PDF")
    
    def _extract_from_doc(self, filepath: str) -> str:
        """Extraire le texte d'un fichier DOC (ancien format Word)"""
        try:
            # Utiliser python-docx2txt ou antiword si disponible
            import docx2txt
            text = docx2txt.process(filepath)
            return self._clean_text(text) if text else ""
        except ImportError:
            raise Exception("Module docx2txt non disponible pour les fichiers .doc")
        except Exception as e:
            raise Exception(f"Erreur lors de la lecture du fichier DOC: {str(e)}")
    
    def _extract_with_ocr(self, filepath: str) -> str:
        """Extraire le texte avec OCR (pour documents scannés)"""
        try:
            import pdf2image
            
            # Convertir PDF en images
            images = pdf2image.convert_from_path(filepath, dpi=300)
            
            full_text = []
            for i, image in enumerate(images):
                # Appliquer OCR sur chaque page
                text = pytesseract.image_to_string(image, lang='fra+eng')
                if text.strip():
                    full_text.append(text)
            
            return '\n'.join(full_text)
            
        except Exception as e:
            raise Exception(f"Erreur OCR: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Nettoyer et normaliser le texte extrait"""
        if not text:
            return ""
        
        # Supprimer les caractères de contrôle
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        
        # Normaliser les espaces
        import re
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Supprimer les lignes vides multiples
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def validate_document(self, filepath: str) -> dict:
        """Valider un document avant traitement"""
        validation_result = {
            'valid': False,
            'errors': [],
            'warnings': [],
            'file_info': {}
        }
        
        try:
            # Vérifier l'existence
            if not os.path.exists(filepath):
                validation_result['errors'].append("Fichier non trouvé")
                return validation_result
            
            # Informations du fichier
            file_size = os.path.getsize(filepath)
            file_extension = os.path.splitext(filepath)[1].lower()
            
            validation_result['file_info'] = {
                'size': file_size,
                'size_mb': round(file_size / 1024 / 1024, 2),
                'extension': file_extension,
                'name': os.path.basename(filepath)
            }
            
            # Vérifier la taille
            if file_size > self.max_file_size:
                validation_result['errors'].append(f"Fichier trop volumineux ({validation_result['file_info']['size_mb']} MB > 50 MB)")
            
            # Vérifier le format
            if file_extension not in self.supported_formats:
                validation_result['errors'].append(f"Format non supporté: {file_extension}")
            
            # Vérifier l'intégrité du fichier
            try:
                if file_extension == '.docx':
                    with zipfile.ZipFile(filepath, 'r') as zip_file:
                        zip_file.testzip()
                elif file_extension == '.pdf':
                    with open(filepath, 'rb') as f:
                        PyPDF2.PdfReader(f)
            except Exception as e:
                validation_result['errors'].append(f"Fichier corrompu: {str(e)}")
            
            # Vérifications de sécurité
            if file_size == 0:
                validation_result['errors'].append("Fichier vide")
            
            # Si pas d'erreurs, marquer comme valide
            if not validation_result['errors']:
                validation_result['valid'] = True
            
        except Exception as e:
            validation_result['errors'].append(f"Erreur de validation: {str(e)}")
        
        return validation_result
    
    def get_document_metadata(self, filepath: str) -> dict:
        """Extraire les métadonnées du document"""
        metadata = {
            'filename': os.path.basename(filepath),
            'size': os.path.getsize(filepath),
            'format': os.path.splitext(filepath)[1].lower(),
            'created': None,
            'modified': None,
            'author': None,
            'title': None,
            'pages': None,
            'word_count': None
        }
        
        try:
            # Métadonnées système
            stat = os.stat(filepath)
            metadata['created'] = stat.st_ctime
            metadata['modified'] = stat.st_mtime
            
            # Métadonnées spécifiques au format
            if metadata['format'] == '.docx':
                metadata.update(self._get_docx_metadata(filepath))
            elif metadata['format'] == '.pdf':
                metadata.update(self._get_pdf_metadata(filepath))
            
        except Exception as e:
            print(f"Erreur lors de l'extraction des métadonnées: {e}")
        
        return metadata
    
    def _get_docx_metadata(self, filepath: str) -> dict:
        """Extraire les métadonnées d'un fichier DOCX"""
        metadata = {}
        
        try:
            doc = docx.Document(filepath)
            
            # Propriétés du document
            props = doc.core_properties
            metadata['author'] = props.author
            metadata['title'] = props.title
            metadata['created'] = props.created
            metadata['modified'] = props.modified
            
            # Compter les pages (approximation)
            metadata['pages'] = len(doc.sections)
            
            # Compter les mots
            word_count = 0
            for paragraph in doc.paragraphs:
                word_count += len(paragraph.text.split())
            metadata['word_count'] = word_count
            
        except Exception as e:
            print(f"Erreur métadonnées DOCX: {e}")
        
        return metadata
    
    def _get_pdf_metadata(self, filepath: str) -> dict:
        """Extraire les métadonnées d'un fichier PDF"""
        metadata = {}
        
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Nombre de pages
                metadata['pages'] = len(pdf_reader.pages)
                
                # Métadonnées du document
                if pdf_reader.metadata:
                    metadata['title'] = pdf_reader.metadata.get('/Title', '')
                    metadata['author'] = pdf_reader.metadata.get('/Author', '')
                    metadata['creator'] = pdf_reader.metadata.get('/Creator', '')
                    metadata['producer'] = pdf_reader.metadata.get('/Producer', '')
                
        except Exception as e:
            print(f"Erreur métadonnées PDF: {e}")
        
        return metadata